# -*- coding: utf-8 -*-
"""
Created on Fri Nov  1 18:33:36 2024

@author: mcall
"""



import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
import spacy
import pandas as pd
import pytesseract
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document as WordDocument
from PyPDF2 import PdfReader
from openpyxl import load_workbook
from transformers import pipeline
from nltk.corpus import wordnet
from pdf2image import convert_from_path

# Suppress TensorFlow oneDNN warning
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'

# Suppress Hugging Face symlink warning on Windows
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# Load SpaCy language model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("SpaCy language model 'en_core_web_sm' is not installed. Run 'python -m spacy download en_core_web_sm' to install it.")

# Initialize Summarization Pipeline with explicit model to avoid default warning
try:
    summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
except Exception as e:
    print("Error initializing summarizer pipeline:", e)

# Synonym Expansion for Keywords
keywords = ["example", "test"]
expanded_keywords = set(keywords)
for word in keywords:
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            expanded_keywords.add(lemma.name())

# Function to analyze text with NLP
def analyze_text_with_nlp(text, keywords):
    doc = nlp(text)
    found_keywords = {}
    
    # Search for exact and expanded keywords
    for keyword in keywords:
        count = len(re.findall(r'\b' + re.escape(keyword) + r'\b', text, flags=re.IGNORECASE))
        if count > 0:
            found_keywords[keyword] = count

    # Use NER to find named entities or relevant phrases
    for ent in doc.ents:
        if ent.label_ in ["ORG", "GPE", "PERSON", "PRODUCT", "EVENT"]:
            if ent.text not in found_keywords:
                found_keywords[ent.text] = 1
            else:
                found_keywords[ent.text] += 1

    return found_keywords

# Function for OCR on image-based PDFs
def ocr_image(image_path):
    return pytesseract.image_to_string(image_path)

# Function to generate a word cloud from found keywords
def generate_wordcloud(found_keywords):
    if found_keywords:
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate_from_frequencies(found_keywords)
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.show()

# Document Processing Functions

# Process Word documents (.docx)
def process_word_doc(file_path):
    doc = WordDocument(file_path)
    text = " ".join([paragraph.text for paragraph in doc.paragraphs])
    return analyze_text_with_nlp(text, expanded_keywords)

# Process Excel files (.xlsx)
def process_excel(file_path):
    workbook = load_workbook(file_path)
    text = ""
    for sheet in workbook.sheetnames:
        worksheet = workbook[sheet]
        for row in worksheet.iter_rows(values_only=True):
            text += " ".join(str(cell) for cell in row if cell)
    return analyze_text_with_nlp(text, expanded_keywords)

# Process PDF files (.pdf)
def process_pdf(file_path):
    reader = PdfReader(file_path)
    text = " ".join(page.extract_text() for page in reader.pages if page.extract_text())
    
    # If text extraction fails, use OCR on images of each PDF page
    if not text:
        images = convert_from_path(file_path)
        text = " ".join(ocr_image(image) for image in images)
    
    return analyze_text_with_nlp(text, expanded_keywords)

# GUI functionality to load a folder and run analysis
def load_folder():
    folder_path = filedialog.askdirectory()
    if not folder_path:
        return

    report = WordDocument()
    report.add_heading("Enhanced Files Analysis Report", 0)
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        found_words = None
        
        # Check file type and process accordingly
        if filename.endswith(".docx"):
            found_words = process_word_doc(file_path)
        elif filename.endswith(".xlsx"):
            found_words = process_excel(file_path)
        elif filename.endswith(".pdf"):
            found_words = process_pdf(file_path)
        else:
            print(f"Skipping unsupported file type: {filename}")
            continue  # Skip unsupported files

        # Add analysis to report if keywords are found
        if found_words:
            report.add_heading(f"Analysis for {filename}", level=1)
            for word, count in found_words.items():
                report.add_paragraph(f"{word}: {count} occurrences")

            # Generate summary if keywords are found
            try:
                summarized_text = summarizer(" ".join(found_words.keys()), max_length=10, min_length=5, do_sample=False)
                report.add_paragraph("Summary: " + summarized_text[0]['summary_text'])
            except Exception as e:
                print("Summarization error:", e)

            # Generate word cloud for visualization
            generate_wordcloud(found_words)
        else:
            report.add_paragraph(f"No keywords or relevant phrases found in {filename}")

    # Save the report in the selected folder
    report.save(os.path.join(folder_path, "Enhanced_Files_Analysis_Report.docx"))
    messagebox.showinfo("Analysis Complete", "Report saved as 'Enhanced_Files_Analysis_Report.docx'")

# Set up GUI
root = tk.Tk()
root.title("Document Analysis Tool")
root.geometry("400x200")

label = tk.Label(root, text="Select a folder to analyze and press Load Folder:")
label.pack(pady=10)

load_button = tk.Button(root, text="Load Folder", command=load_folder)
load_button.pack(pady=10)

root.mainloop()
