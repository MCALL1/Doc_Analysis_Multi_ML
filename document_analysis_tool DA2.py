# -*- coding: utf-8 -*-
"""
Created on Sat Nov  2 11:59:59 2024


@author: mcall
"""

 """    The project, Legal Document Analyzer with Summarization and Keyword Extraction, integrates several advanced technologies to enhance the analysis of legal documents:

   Natural Language Processing (NLP):
        SpaCy: Utilized for tokenization, named entity recognition (NER), and part-of-speech tagging, enabling the extraction of entities such as organizations, locations, and individuals from text.
        NLTK (Natural Language Toolkit): Employed for synonym expansion using WordNet, broadening keyword searches to include semantically related terms.

   Machine Learning for Text Summarization:
        Hugging Face Transformers: Incorporates the sshleifer/distilbart-cnn-12-6 model, a distilled version of BART (Bidirectional and Auto-Regressive Transformer), to generate concise summaries of lengthy legal documents.

   Optical Character Recognition (OCR):
        pytesseract: Applies OCR to extract text from image-based PDFs, ensuring that non-text PDFs are also analyzable.

   Data Visualization:
        WordCloud: Generates visual representations of keyword frequencies, aiding in the quick identification of prevalent terms within documents.
        Matplotlib: Facilitates the display of word clouds and other potential visualizations.

   Graphical User Interface (GUI):
        Tkinter: Provides a user-friendly interface for selecting folders, initiating analysis, and displaying results, making the tool accessible to users without programming expertise.

   Document Processing:
        python-docx: Handles the reading and writing of Word documents, enabling the extraction and reporting of text.
        openpyxl: Manages Excel files, allowing for the extraction of textual data from spreadsheets.
        pdfplumber: Extracts text from PDFs, offering improved accuracy over some other PDF text extraction methods.

   Date and Monetary Amount Extraction:
        dateparser: Parses and extracts dates from text, accommodating various date formats.
        Regular Expressions (regex): Identifies monetary amounts within text, supporting multiple currency symbols and formats.

   By integrating these technologies, the project offers a comprehensive solution for analyzing legal documents, extracting pertinent information, and providing concise summaries, thereby streamlining the review process for legal professionals.
#                          Summary of functionality
#1. Document Type Handling:

    # The tool processes various document types: Word documents (.docx), Excel spreadsheets (.xlsx), and PDFs (.pdf).
    # It can handle both text-based and image-based PDFs, using OCR to extract text from scanned documents when necessary.

#2. Keyword Analysis:

    # The tool analyzes documents for specific legal-related keywords (e.g., "contract," "agreement," "payment").
    # It uses SpaCy for Named Entity Recognition (NER) to detect named entities such as organizations, people, products, and events.
    # Synonym Expansion: With NLTK’s WordNet, it expands keywords by including synonyms, making keyword search more comprehensive.

#3. Date and Monetary Amount Extraction:

    # Dates: Uses dateparser to extract dates, accommodating various date formats across documents.
    # Monetary Amounts: Uses regular expressions to identify and extract currency values, capturing amounts with symbols like $, £, and €.

#4. Text Summarization:

    # Transformers Pipeline: Generates a concise summary of each document, providing a quick overview of the contents.
    # The Hugging Face transformers library is used to produce summaries based on extracted keywords.

#5. Visualization with Word Cloud:

    # For each document, the tool creates a word cloud to visualize the frequency of keywords, giving a quick snapshot of the document’s focus.

#6. Report Generation:

    # Comprehensive Report: For each processed document, the tool generates a report that includes:
        # The frequency of keywords.
        # Extracted dates and monetary amounts.
        # A summary of the document’s key content.
    # The report is saved as a Word document in the selected folder, making it easy to review and share insights.

#7. User-Friendly GUI:

    # A simple GUI built with Tkinter allows users to select a folder, initiate analysis, and receive a notification when the report is complete.
   
#8. Dynamic Summarization:
      #  Introduced a new dynamic_summarization function that adjusts max_length and min_length for the summarizer based on the input text length.
      #  This avoids warnings and ensures efficient summarization for different document lengths.

#9. Efficient Report Generation:
      #  Updated the report generation to use the dynamic_summarization function for summaries, making the summarization process responsive to shorter or longer texts.

 # These changes make the summarization more adaptive, reduce unnecessary warnings, and improve overall efficiency in handling various document types and lengths.
 # Use Cases:

    # 1. Legal Document Analysis: The tool is ideal for processing contracts, agreements, invoices, and other legal documents, extracting relevant information like dates, monetary values, and key legal terms.
    # 2. Compliance and Auditing: Useful for identifying important compliance dates and payment terms within financial and regulatory documents.
    # 3. Summarization and Quick Insights: Helps users quickly review large volumes of documents by generating summaries and visual keyword maps.

 # This comprehensive tool provides both detailed insights and a streamlined overview of each document, making it an efficient solution for legal, financial, and compliance-related text analysis tasks.

#                                End Of Summary 

import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
import spacy
import pandas as pd
import pytesseract
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from docx import Document as WordDocument
from PyPDF2 import PdfReader
import pdfplumber
from openpyxl import load_workbook
from transformers import pipeline
import dateparser
from pdf2image import convert_from_path
from nltk.corpus import wordnet
import nltk

# Ensure NLTK wordnet data is downloaded
nltk.download('wordnet')

# Suppress TensorFlow oneDNN warning
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# Load SpaCy language model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("SpaCy language model 'en_core_web_sm' is not installed. Run 'python -m spacy download en_core_web_sm' to install it.")


#                   Machine Learning Component: Text Summarization
# 
# This script uses a machine learning model from the Hugging Face Transformers library 
# to generate summaries for each document. The model used is "sshleifer/distilbart-cnn-12-6", 
# a smaller, optimized version of the BART model designed for efficient summarization tasks.
#
# About the Model:
#   The model "sshleifer/distilbart-cnn-12-6" is a transformer-based sequence-to-sequence model 
#   fine-tuned for abstractive text summarization. It is trained to generate concise summaries 
#   that capture the main points of longer text inputs.
#   BART (Bidirectional and Auto-Regressive Transformer) is effective for tasks like summarization 
#   and translation due to its ability to encode input text and generate new text (a summary) as output.
#
# Model Initialization:
#   We initialize the summarizer pipeline using the specific model name. By setting the model explicitly, 
#   we avoid loading a default model and ensure consistency in output across runs.
#   If the text length of the input is very short, the script dynamically adjusts `max_length` and 
#   min_length` to optimize summary generation and prevent unnecessary warnings.
#
# Purpose:
#   The summarizer adds value by automatically generating brief summaries for each document, allowing 
#   quick insight into the document’s content without manually reading through it.
#   This approach leverages pre-trained machine learning models to deliver intelligent summarization 
#   without the need for custom model training, making the tool both efficient and easy to use.
#                                      End of Summary

 #Initialize Summarization Pipeline with explicit model to avoid default warning
try:
    summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
except Exception as e:
    print("Error initializing summarizer pipeline:", e)

# Synonym Expansion for Keywords (Including Legal Terms)
keywords = ["contract", "agreement", "payment", "compliance", "fine", "lease", "memo" "contract", "agreement", "payment", "compliance", "fine", "lease", "memo", "settlement", "liability", "warranty", 
"litigation", "dispute", "arbitration", "breach", "damages", "termination", "compensation", "indemnity", 
"disclosure", "confidentiality", "jurisdiction", "arbitrator", "liquidated", "damages", "provision", "statute", 
"enforceability", "jurisprudence", "license", "penalty", "waiver", "void", "voidable", "binding", "negotiation", 
"acceptance", "offer", "counteroffer", "party", "performance", "execution", "rescission", "tort", "negligence", 
"strict liability", "remedy", "mitigation", "equity", "legal capacity", "consideration", "intent", "clause", 
"estoppel", "fiduciary", "representation", "guarantee", "promissory", "estoppel", "good faith", "rescission", 
"novation", "assignor", "assignee", "novation", "stipulation", "precedent", "recession", "novation", "subrogation", 
"forbearance", "estoppel", "nondisclosure", "prosecution", "defendant", "plaintiff", "attorney", "legal counsel", 
"injunction", "subpoena", "testimony", "verdict", "sentence", "indictment", "appeal", "juror", "jury", "damages", 
"burden of proof", "witness", "deposition", "cross-examination", "discovery", "pleading", "complaint", "petition", 
"motion", "hearing", "summons", "dismissal", "discharge", "acquittal", "settlement", "hearing", "verdict", 
"contempt", "felony", "misdemeanor", "bail", "parole", "probation", "restitution", "legal precedent", 
"emancipation", "judgment", "executory", "parol evidence", "rescission", "constructive", "vicarious liability", 
"due diligence", "employer", "employee", "employment agreement", "severance", "non-compete", "non-solicitation", 
"intellectual property", "trademark", "patent", "copyright", "licensing", "infringement", "trade secret", 
"commercial law", "consumer protection", "product liability", "agency", "principal", "third-party", "personal injury", 
"real estate", "property", "title", "deed", "mortgage", "lien", "foreclosure", "easement", "zoning", "leasehold", 
"tenant", "landlord", "tenancy", "assignment", "sublease", "security deposit", "eviction", "mediation", "court order", 
"summary judgment", "motion to dismiss", "amicus brief", "appellate", "civil procedure", "criminal procedure", 
"burden of proof", "beyond a reasonable doubt", "preponderance of the evidence", "judicial review", "habeas corpus", 
"double jeopardy", "self-incrimination", "precedent", "stare decisis", "doctrine", "mens rea", "actus reus", 
"probable cause", "search warrant", "exclusionary rule", "bail bond", "exculpatory", "custody", "arraignment", 
"plea bargain", "beyond reasonable doubt", "statute of limitations", "probate", "executor", "beneficiary", 
"inheritance", "trustee", "settlor", "estate", "tax", "exemption", "fiduciary duty", "bankruptcy", "debtor", "creditor"]

expanded_keywords = set(keywords)

# Expand keywords with synonyms using NLTK's wordnet
for word in keywords:
    for syn in wordnet.synsets(word):
        for lemma in syn.lemmas():
            expanded_keywords.add(lemma.name())

# Function to analyze text with NLP
def analyze_text_with_nlp(text, keywords):
    doc = nlp(text)
    found_keywords = {}
    
    # Search for exact and expanded keywords
    for keyword in keywords:
        count = len(re.findall(r'\b' + re.escape(keyword) + r'\b', text, flags=re.IGNORECASE))
        if count > 0:
            found_keywords[keyword] = count

    # Use NER to find named entities or relevant phrases
    for ent in doc.ents:
        if ent.label_ in ["ORG", "GPE", "PERSON", "PRODUCT", "EVENT","Supreme Court", "Federal Reserve", "NASDAQ", "Wall Street Journal", "Goldman Sachs", "Department of Justice", "International Monetary Fund", "SEC", "World Bank", "Bloomberg",
 "New York Stock Exchange", "KPMG", "Deloitte", "PwC", "Federal Trade Commission", "Harvard Law Review", "Oxford University Press", "Cambridge University Press", "Blackstone Chambers", "Legal Aid Society",
 "New York", "Washington, D.C.", "London", "Tokyo", "Zurich", "Frankfurt", "Hong Kong", "Singapore", "Paris", "Toronto",
 "Securities Act", "Bankruptcy Code", "Sarbanes-Oxley Act", "Dodd-Frank Act", "Uniform Commercial Code", "Fair Credit Reporting Act", "Sherman Antitrust Act", "Foreign Corrupt Practices Act", "Freedom of Information Act", "Clean Air Act",
 "Moody's", "Fitch Ratings", "Standard & Poor's", "Lehman Brothers", "Bank of America", "JP Morgan Chase", "Morgan Stanley", "UBS", "HSBC", "ING Group",
 "Chicago", "Frankfurt", "Beijing", "Dubai", "Geneva", "Vienna", "Brussels", "Luxembourg", "Los Angeles", "Boston",
 "Lending Agreement", "Insurance Policy", "Service Contract", "Lease Agreement", "Power of Attorney", "Employment Contract", "Settlement Agreement", "Non-Disclosure Agreement", "Mortgage Contract", "Bill of Sale",
 "BlackRock", "Vanguard", "Citigroup", "Royal Bank of Canada", "Wells Fargo", "Credit Suisse", "Deutsche Bank", "American Express", "Visa", "Mastercard",
 "Shanghai", "Sydney", "San Francisco", "Johannesburg", "Montreal", "Milan", "Brussels", "Cape Town", "Mumbai", "Seoul",
 "Shareholder Agreement", "Purchase Agreement", "Divorce Decree", "Arbitration Award", "Court Order", "Affidavit", "Probate Document", "Summons", "Subpoena", "Contract",
 "American Bar Association", "Federal Bureau of Investigation", "National Association of Securities Dealers", "World Economic Forum", "Association of American Law Schools", "Human Rights Watch", "Inter-American Development Bank", "Transparency International", "UNICEF", "International Criminal Court",
 "Buenos Aires", "Riyadh", "Bangkok", "Lagos", "Rio de Janeiro", "Helsinki", "Santiago", "Ankara", "Jakarta", "Kuala Lumpur",
 "Patent", "Trademark", "Copyright", "Lease", "Deed", "Warranty", "Lien", "Judgment", "Plea Bargain", "Bankruptcy",
 "Accenture", "McKinsey & Company", "Boston Consulting Group", "Ernst & Young", "Morgan Stanley", "Barclays", "Credit Agricole", "Societe Generale", "BNP Paribas", "ICBC",
 "Ottawa", "Copenhagen", "Manila", "Nairobi", "Amsterdam", "Prague", "Stockholm", "Mexico City", "Hanoi", "Warsaw",
 "Class Action Lawsuit", "Restraining Order", "Injunction", "Deposition", "Mediation Agreement", "Custody Agreement", "Property Deed", "Non-Compete Agreement", "Arbitration Agreement", "Severance Package",
 "Internal Revenue Service", "European Central Bank", "Asian Development Bank", "Bank for International Settlements", "American Arbitration Association", "World Intellectual Property Organization", "International Bar Association", "National Labor Relations Board", "Consumer Financial Protection Bureau", "Office of the Comptroller of the Currency",
 "Bail Bond", "Arraignment", "Litigation", "Discovery", "Indemnity Clause", "Force Majeure Clause", "Confidentiality Agreement", "Limited Partnership Agreement", "Corporate Bylaws", "Employment Offer Letter",
 "Merrill Lynch", "T. Rowe Price", "Rothschild", "Fidelity Investments", "Charles Schwab", "UBS Wealth Management", "Allianz", "Manulife", "Prudential Financial", "AXA",
 "Standard Chartered", "Rabobank", "Mizuho Bank", "Nordea", "Danske Bank", "Santander", "Unicredit", "Sumitomo Mitsui Banking Corporation", "Bank of China", "Royal Bank of Scotland",
 "Title Insurance", "Real Estate Deed", "Shareholder Proxy", "Asset Purchase Agreement", "Settlement Check", "Proof of Funds", "Escrow Agreement", "Indemnification Agreement", "Business License", "Power of Attorney Document",
 "Intergovernmental Panel on Climate Change", "International Finance Corporation", "European Union Agency for Fundamental Rights", "Amnesty International", "Commonwealth Bank", "National Credit Union Administration", "Federal Deposit Insurance Corporation", "Trade Adjustment Assistance Program", "Financial Action Task Force", "The Hague Conference on Private International Law",
 "Equal Employment Opportunity Commission", "Health and Safety Executive", "Labor Board", "Anti-Money Laundering Act", "Gramm-Leach-Bliley Act", "Employment Rights Act", "Corporate Governance Code", "Private Securities Litigation Reform Act", "Whistleblower Protection Act", "Securities Exchange Act",
 "Reserve Bank of India", "People's Bank of China", "Bank of Japan", "Swiss National Bank", "Bank of England", "European Investment Bank", "Federal Deposit Insurance Corporation", "Export-Import Bank of the United States", "Asian Infrastructure Investment Bank", "African Development Bank"]:

            if ent.text not in found_keywords:
                found_keywords[ent.text] = 1
            else:
                found_keywords[ent.text] += 1

    return found_keywords

# Function to extract dates and monetary amounts
def extract_info(text):
    dates = []
    amounts = []
    
    # Extract dates using dateparser
    for line in text.splitlines():
        parsed_date = dateparser.parse(line)
        if parsed_date:
            dates.append(parsed_date)

    # Use regex to find monetary values
    amount_pattern = r'[\$\£\€]?\s?\d+(?:,\d{3})*(?:\.\d{2})?\s?(USD|GBP|EUR)?'
    amounts.extend(re.findall(amount_pattern, text))
    
    return dates, amounts

# Function for OCR on image-based PDFs
def ocr_image(image):
    """Extracts text from an image using OCR."""
    return pytesseract.image_to_string(image)

# Function to generate a word cloud from found keywords
def generate_wordcloud(found_keywords):
    if found_keywords:
        wordcloud = WordCloud(width=800, height=400, background_color='white').generate_from_frequencies(found_keywords)
        plt.figure(figsize=(10, 5))
        plt.imshow(wordcloud, interpolation='bilinear')
        plt.axis("off")
        plt.show()

# Dynamic Summarization Function
def dynamic_summarization(text):
    """Adjusts max_length and min_length based on the input text length for summarization."""
    input_length = len(text.split())

    # Set max and min length dynamically based on input length
    if input_length < 20:
        max_length = 5
        min_length = 3
    elif input_length < 50:
        max_length = 10
        min_length = 5
    else:
        max_length = 50
        min_length = 25

    try:
        summarized_text = summarizer(text, max_length=max_length, min_length=min_length, do_sample=False)
        return summarized_text[0]['summary_text']
    except Exception as e:
        print("Summarization error:", e)
        return "Summary could not be generated."

# Document Processing Functions

# Process Word documents (.docx)
def process_word_doc(file_path):
    doc = WordDocument(file_path)
    text = " ".join([paragraph.text for paragraph in doc.paragraphs])
    dates, amounts = extract_info(text)
    print(f"Extracted Dates: {dates}")
    print(f"Extracted Monetary Amounts: {amounts}")
    return analyze_text_with_nlp(text, expanded_keywords), dates, amounts

# Process Excel files (.xlsx)
def process_excel(file_path):
    workbook = load_workbook(file_path)
    text = ""
    for sheet in workbook.sheetnames:
        worksheet = workbook[sheet]
        for row in worksheet.iter_rows(values_only=True):
            text += " ".join(str(cell) for cell in row if cell)
    dates, amounts = extract_info(text)
    print(f"Extracted Dates: {dates}")
    print(f"Extracted Monetary Amounts: {amounts}")
    return analyze_text_with_nlp(text, expanded_keywords), dates, amounts

# Process PDF files (.pdf)
def process_pdf(file_path):
    text = ""
    # Use pdfplumber for improved text extraction
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    
    # If text extraction fails (indicating an image-based PDF), use OCR
    if not text.strip():
        images = convert_from_path(file_path)
        text = " ".join(ocr_image(image) for image in images)
    
    dates, amounts = extract_info(text)
    print(f"Extracted Dates: {dates}")
    print(f"Extracted Monetary Amounts: {amounts}")
    return analyze_text_with_nlp(text, expanded_keywords), dates, amounts

# GUI functionality to load a folder and run analysis
def load_folder():
    folder_path = filedialog.askdirectory()
    if not folder_path:
        return

    report = WordDocument()
    report.add_heading("Enhanced Files Analysis Report", 0)
    
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        found_words, dates, amounts = None, None, None
        
        # Check file type and process accordingly
        if filename.endswith(".docx"):
            found_words, dates, amounts = process_word_doc(file_path)
        elif filename.endswith(".xlsx"):
            found_words, dates, amounts = process_excel(file_path)
        elif filename.endswith(".pdf"):
            found_words, dates, amounts = process_pdf(file_path)
        else:
            print(f"Skipping unsupported file type: {filename}")
            continue  # Skip unsupported files

        # Add analysis to report if keywords are found
        if found_words:
            report.add_heading(f"Analysis for {filename}", level=1)
            for word, count in found_words.items():
                report.add_paragraph(f"{word}: {count} occurrences")

            # Add extracted dates and monetary amounts
            if dates:
                report.add_paragraph("Dates: " + ", ".join(str(date) for date in dates))
            if amounts:
                report.add_paragraph("Monetary Amounts: " + ", ".join(amounts))

            # Generate summary using dynamic summarization
            summary_text = dynamic_summarization(" ".join(found_words.keys()))
            report.add_paragraph("Summary: " + summary_text)

            # Generate word cloud for visualization
            generate_wordcloud(found_words)
        else:
            report.add_paragraph(f"No keywords or relevant phrases found in {filename}")

    # Save the report in the selected folder
    report.save(os.path.join(folder_path, "Enhanced_Files_Analysis_Report.docx"))
    messagebox.showinfo("Analysis Complete", "Report saved as 'Enhanced_Files_Analysis_Report.docx'")

# Set up GUI
root = tk.Tk()
root.title("Document Analysis Tool")
root.geometry("400x200")

label = tk.Label(root, text="Select a folder to analyze and press Load Folder:")
label.pack(pady=10)

load_button = tk.Button(root, text="Load Folder", command=load_folder)
load_button.pack(pady=10)

root.mainloop()



